# -*- coding: utf-8 -*-
"""
提案腳本（含分鏡）docx 產生器 — 對標「提案腳本架構（範本）.docx」

沿用綠芽既有 docx 規格：A4 直式、邊距 0.5"、表格總寬 10449 dxa、標題 #0176BB。
照片區留空（參考圖片欄空白），由設計/後製後續貼圖。

用法：
    from build_script_doc import build_script_doc
    build_script_doc(client_name="品香茶葉", info={...}, scenes=[{...}], out_path="...docx")

info schema：
    core      核心訊息（主題框架 + 一句標語 + 50字說明）
    story     故事大綱（3–5 句）
    mood      風格/情緒（2–3 個形容詞）
    audience  目標觀眾
    spec      dict(片長, 規格, 語版, 交付, 演員)   ← 一律來自「影片需求單」，不是範本預設值
    shoot     -拍攝手法- 段落文字
    shoot_ref 拍攝手法參考影片（URL 或描述，選填）
    fx        -特效風格- 段落文字
    fx_ref    特效風格參考影片（選填）
    board_title  分鏡標題（放在「腳本分鏡【】」括號內，如客戶名稱或片名）

scene schema（每段一個 dict）：
    seg   段落（如「1. 核心開場 0:00–0:15」）
    card  字卡（沒有就空字串）
    img   參考圖片（一般留空，給後續貼圖；可填示意文字）
    visual 影像說明（@ 開頭、每段 3–5 個鏡頭指示，可用 \\n 斷行）
    vo    旁白/備註（語言依需求單語版）
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

BLUE = RGBColor(0x01, 0x76, 0xBB)
GREY = RGBColor(0xAE, 0xAA, 0xAA)
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
EA = "PingFang TC"

# 5 欄寬（dxa），總和 10449，沿用 script-writer 既有規格
COLW = [1119, 1890, 2834, 2409, 2197]
HEADERS = ["段落", "字卡", "參考圖片", "影像說明", "旁白/備註"]


def _set_run(run, size=10, bold=False, color=BLACK, ea=EA):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = ea
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), ea)


def _para(doc, label=None, text="", size=10, bold=False, color=BLACK,
          label_bold=False, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if label:
        _set_run(p.add_run(label), size=size, bold=label_bold or bold, color=color)
    if text:
        _set_run(p.add_run(text), size=size, bold=bold, color=color)
    return p


def _set_cell_width(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):     # 先移除自動產生的，否則 Word 只讀第一個等寬值
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(dxa)); tcW.set(qn('w:type'), 'dxa')
    tcPr.insert(0, tcW)


def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _cell_text(cell, text, size=9, bold=False, color=BLACK):
    cell.text = ""
    first = True
    for line in str(text).split("\n"):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        _set_run(p.add_run(line), size=size, bold=bold, color=color)


def _borders_inner_only(table):
    """無外框，僅內部分隔線 single #AEAAAA"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:val'), 'none'); borders.append(e)
    for edge in ('insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:color'), 'AEAAAA'); borders.append(e)
    tblPr.append(borders)


def build_script_doc(client_name, info, scenes, out_path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)        # A4 直式
    sec.left_margin = sec.right_margin = Inches(0.5)
    sec.top_margin = sec.bottom_margin = Inches(0.5)

    # 標題
    _para(doc, label="企業形象影片", size=14, bold=True, color=BLUE, space_after=6)

    # 腳本資訊
    _para(doc, "核心訊息：", info.get("core", ""))
    _para(doc, "故事大綱：", info.get("story", ""))
    _para(doc, "風格/情緒：", info.get("mood", ""))
    _para(doc, "目標觀眾：", info.get("audience", ""), space_after=8)

    # 影片規格（一律來自需求單）
    _para(doc, label="影片規格", size=11, bold=True, color=BLACK, space_after=3)
    spec = info.get("spec", {})
    for k in ("片長", "規格", "語版", "交付", "演員"):
        if spec.get(k):
            _para(doc, f"{k}：", spec[k])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # 影像風格
    _para(doc, label="影像風格", size=11, bold=True, color=BLACK, space_after=3)
    _para(doc, label="-拍攝手法-", size=10, bold=True, space_after=1)
    _para(doc, text=info.get("shoot", ""))
    if info.get("shoot_ref"):
        _para(doc, label="參考影片：", text=info["shoot_ref"], label_bold=True, color=GREY)
    _para(doc, label="-特效風格-", size=10, bold=True, space_after=1)
    _para(doc, text=info.get("fx", ""))
    if info.get("fx_ref"):
        _para(doc, label="參考影片：", text=info["fx_ref"], label_bold=True, color=GREY)

    # 分鏡標題
    bt = info.get("board_title", client_name)
    length = info.get("spec", {}).get("片長", "")
    _para(doc, label=f"腳本分鏡【{bt}】{length}", size=10, bold=True, space_after=1)
    _para(doc, text="*圖片作為參考構圖", size=8, color=GREY, space_after=4)

    # 分鏡表
    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    table.allow_autofit = False
    _borders_inner_only(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(HEADERS):
        _cell_text(hdr[i], h, size=10, bold=True, color=WHITE)
        _shade(hdr[i], "0176BB")
        _set_cell_width(hdr[i], COLW[i])
    for sc in scenes:
        row = table.add_row().cells
        vals = [sc.get("seg", ""), sc.get("card", ""), sc.get("img", ""),
                sc.get("visual", ""), sc.get("vo", "")]
        for i, v in enumerate(vals):
            _cell_text(row[i], v, size=9)
            _set_cell_width(row[i], COLW[i])

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    # 煙霧測試：確認 docx 能正常產生（內容為佔位示意，非真實腳本）
    info = dict(
        core="主題框架：示意。一句標語：示意。說明：示意說明段落。",
        story="示意故事大綱三到五句。", mood="沉穩、自然、國際", audience="海外進口商/通路採購",
        spec={"片長": "1–2 分鐘", "規格": "比例 16:9，4K（2160p），輸出 MP4",
              "語版": "英文旁白／英文字卡／英文字幕", "交付": "有字幕版、無字幕版",
              "演員": "依需求單"},
        shoot="示意拍攝手法段落。", shoot_ref="https://example.com/ref1",
        fx="示意特效風格段落。", fx_ref="",
        board_title="煙霧測試")
    scenes = [
        {"seg": "1. 核心開場 0:00–0:15", "card": "MADE IN TAIWAN", "img": "",
         "visual": "@晨霧茶山空拍\n@採茶人剪影", "vo": "Some tea begins with the land."},
        {"seg": "7. 品牌精神 1:50–2:00", "card": "Slogan + LOGO", "img": "",
         "visual": "@成品上棧板\n@LOGO 動畫收尾", "vo": "Pin Xiang — Tea That Could Only Be Taiwan."},
    ]
    out = "/tmp/_newscript_smoke.docx"
    print("已產生：", build_script_doc("煙霧測試", info, scenes, out))
